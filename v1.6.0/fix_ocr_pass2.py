#!/usr/bin/env python3
"""fix_ocr_pass2.py — Second-pass OCR cleanup for iqt_v1_6_0.docx.

Fixes ~30 remaining errors across 6 categories:
1. Figure renumbering (text replacement)
2. Heading style fixes (style property change)
3. Figure caption style fix (style + remove empty run)
4. Inline newline removal (remove w:br elements)
5. Orphan word merges (backward, deletes paragraphs)
6. Mid-sentence paragraph merges (backward, deletes paragraphs)
"""

import copy
import sys
from docx import Document
from lxml import etree

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

DOCX = "iqt_v1_6_0.docx"


def find_para(paragraphs, prefix, start=0):
    """Return the index of the first paragraph whose text starts with `prefix`."""
    for i in range(start, len(paragraphs)):
        if paragraphs[i].text.lstrip("\n").lstrip().startswith(prefix):
            return i
    raise ValueError(f"Paragraph not found starting from {start}: {prefix!r}")


def find_para_endswith(paragraphs, suffix, start=0):
    """Return the index of the first paragraph whose text ends with `suffix`."""
    for i in range(start, len(paragraphs)):
        if paragraphs[i].text.rstrip().endswith(suffix):
            return i
    raise ValueError(f"Paragraph ending with {suffix!r} not found from {start}")


def merge_para_into(doc, idx_target, idx_source, space=""):
    """Append all runs from source paragraph into target, then delete source.

    `space` is inserted between the last target run and the first source run
    (typically " " or "").
    """
    target = doc.paragraphs[idx_target]
    source = doc.paragraphs[idx_source]

    # If we need a space, append it to the last run of target (or create one)
    if space and target.runs:
        target.runs[-1].text += space
    elif space:
        r = target.add_run(space)

    # Copy each run from source, preserving formatting
    for run in source.runs:
        new_run = target.add_run(run.text)
        # Copy run properties XML to preserve all formatting
        if run._element.find("w:rPr", NSMAP) is not None:
            new_rpr = copy.deepcopy(run._element.find("w:rPr", NSMAP))
            existing = new_run._element.find("w:rPr", NSMAP)
            if existing is not None:
                new_run._element.remove(existing)
            new_run._element.insert(0, new_rpr)

    # Delete source paragraph from the document body
    source._element.getparent().remove(source._element)


def remove_breaks(para):
    """Remove all w:br elements from a paragraph."""
    brs = para._element.findall(".//w:br", NSMAP)
    count = 0
    for br in brs:
        br.getparent().remove(br)
        count += 1
    return count


def remove_empty_leading_run(para):
    """Remove the first run if it is empty (text is '' or whitespace-only)."""
    if para.runs and para.runs[0].text.strip() == "":
        para.runs[0]._element.getparent().remove(para.runs[0]._element)
        return True
    return False


def main():
    doc = Document(DOCX)
    paras = doc.paragraphs
    initial_count = len(paras)
    print(f"Loaded {DOCX}: {initial_count} paragraphs")

    # ── 1. Figure renumbering ────────────────────────────────────────────
    print("\n─── 1. Figure renumbering ───")

    renames = [
        ("Figure 2: Nested Causal Diamonds", "Figure 1: Nested Causal Diamonds"),
        ("Figure 5: Three Layers", "Figure 2: Three Layers"),
    ]
    for old, new in renames:
        idx = find_para(paras, old)
        paras[idx].runs[0].text = paras[idx].runs[0].text.replace(old, new)
        # Also check if text spans multiple runs
        full = paras[idx].text
        if old in full:
            # Replace in the full text by finding which run has it
            for r in paras[idx].runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
        print(f"  [{idx}] {old!r} -> {new!r}")

    # Figure 6: Thought-Experiment -> Figure 8 (the one at ~548)
    # Must find the *second* "Figure 6:" since the first is Protocol 2
    idx_fig6a = find_para(paras, "Figure 6:")
    idx_fig6b = find_para(paras, "Figure 6:", start=idx_fig6a + 1)
    for r in paras[idx_fig6b].runs:
        if "Figure 6:" in r.text:
            r.text = r.text.replace("Figure 6:", "Figure 8:")
            break
    print(f"  [{idx_fig6b}] 'Figure 6: Thought-Experiment' -> 'Figure 8:'")

    # ── 2. Heading style fixes ───────────────────────────────────────────
    print("\n─── 2. Heading style fixes ───")

    # "5.4 Supplementary: Split-Brain" -> Heading 2, bold
    idx_h54 = find_para(paras, "5.4 Supplementary: Split-Brain")
    paras[idx_h54].style = doc.styles["Heading 2"]
    for r in paras[idx_h54].runs:
        r.bold = True
    print(f"  [{idx_h54}] set to Heading 2 + bold")

    # "Interpreting Section 6: Where" -> Heading 3 (already bold)
    idx_interp = find_para(paras, "Interpreting Section 6: Where")
    paras[idx_interp].style = doc.styles["Heading 3"]
    print(f"  [{idx_interp}] set to Heading 3")

    # ── 3. Figure caption style fix ──────────────────────────────────────
    print("\n─── 3. Figure caption style fix ───")

    # Find the "Protocol 3: Predicted" paragraph (starts with space)
    idx_proto3 = find_para(paras, "Protocol 3: Predicted")
    paras[idx_proto3].style = doc.styles["Heading 4"]
    remove_empty_leading_run(paras[idx_proto3])
    # Also strip the leading space from first run
    if paras[idx_proto3].runs and paras[idx_proto3].runs[0].text.startswith(" "):
        paras[idx_proto3].runs[0].text = paras[idx_proto3].runs[0].text.lstrip()
    print(f"  [{idx_proto3}] set to Heading 4, cleaned leading space/run")

    # ── 4. Inline newline removal ────────────────────────────────────────
    print("\n─── 4. Inline newline removal ───")

    # Paragraph starting with "Isotony."
    idx_isotony = find_para(paras, "Isotony.")
    n = remove_breaks(paras[idx_isotony])
    # Clean the \n from the run text too
    for r in paras[idx_isotony].runs:
        if "\n" in r.text:
            r.text = r.text.replace("\n", "")
    print(f"  [{idx_isotony}] removed {n} break(s) from Isotony paragraph")

    # Paragraph with leading newline before "Figure 6: Protocol 2"
    # This is the "\nFigure 6:" paragraph — find by the break
    idx_fig6_cap = find_para(paras, "Figure 6: Protocol 2")
    n = remove_breaks(paras[idx_fig6_cap])
    # Clean the \n from run text
    for r in paras[idx_fig6_cap].runs:
        if "\n" in r.text:
            r.text = r.text.replace("\n", "")
    # Remove empty first run if present
    remove_empty_leading_run(paras[idx_fig6_cap])
    print(f"  [{idx_fig6_cap}] removed {n} break(s), cleaned newline text")

    # ── 5. Orphan word merges (backward) ─────────────────────────────────
    print("\n─── 5. Orphan word merges ───")

    # Refresh paragraph list (indices may shift but we use content matching)
    paras = doc.paragraphs

    # "broadly." -> merge into "...self-thread framework"
    idx_broadly = find_para(paras, "broadly.")
    idx_selfthread = find_para_endswith(paras, "self-thread framework")
    print(f"  Merging [{idx_broadly}] 'broadly.' into [{idx_selfthread}]")
    merge_para_into(doc, idx_selfthread, idx_broadly, space=" ")

    # Refresh
    paras = doc.paragraphs

    # "dynamics." -> merge into "...redundancy-dominated"
    idx_dynamics = find_para(paras, "dynamics.")
    idx_redundancy = find_para_endswith(paras, "redundancy-dominated")
    print(f"  Merging [{idx_dynamics}] 'dynamics.' into [{idx_redundancy}]")
    merge_para_into(doc, idx_redundancy, idx_dynamics, space=" ")

    # Refresh
    paras = doc.paragraphs

    # "1: " -> merge into paragraph ending with "embedding:"
    idx_1colon = find_para(paras, "1: ")
    idx_embedding = find_para_endswith(paras, "embedding:")
    print(f"  Merging [{idx_1colon}] '1: ' into [{idx_embedding}]")
    merge_para_into(doc, idx_embedding, idx_1colon, space="\n")

    # ── 6. Mid-sentence paragraph merges (backward by index) ─────────────
    print("\n─── 6. Mid-sentence paragraph merges ───")

    # Build pairs: (first_para_text_end, second_para_text_start, space)
    # We search by content to be robust; merge backward (highest index first)
    merge_specs = [
        # (end-of-first, start-of-second, space_to_insert)
        ("tests in the actual", "experiment: if a non-psychedelic", " "),
        ("Synergy analysis\u2014", "does the overlap zone", ""),
        ("legitimate tests of a theory", "formulated in the language", " "),
        ("quality exists universally and", "experience (Level 1)", " "),
        ("effective subalgebra (\u00a72.6), the", "mutual information", " "),
        ("abstract AQFT and neural", "measurements. The effective", " "),
        ("criterion that could", "trivialize phenomenological", " "),
        ('The term \u201cIntrinsic Quality', '\u201d is chosen deliberately', ""),
        ("Intrinsic Quality", "is the physical state", " "),
    ]

    # Resolve all indices first, then sort descending and merge
    resolved = []
    for end_text, start_text, space in merge_specs:
        paras = doc.paragraphs  # refresh
        idx_first = find_para_endswith(paras, end_text)
        idx_second = find_para(paras, start_text, start=idx_first + 1)
        if idx_second != idx_first + 1:
            print(f"  WARNING: {start_text!r} is at [{idx_second}], not adjacent to [{idx_first}]")
        resolved.append((idx_first, idx_second, space, end_text, start_text))

    # Sort by second paragraph index DESCENDING so deletions don't shift
    resolved.sort(key=lambda x: x[1], reverse=True)

    for idx_first, idx_second, space, end_text, start_text in resolved:
        paras = doc.paragraphs
        # Re-resolve since we're going backward; re-find by content
        idx_first = find_para_endswith(paras, end_text)
        idx_second = find_para(paras, start_text, start=idx_first)
        print(f"  Merging [{idx_second}] '{start_text[:40]}' into [{idx_first}]")
        merge_para_into(doc, idx_first, idx_second, space=space)

    # ── Save ─────────────────────────────────────────────────────────────
    paras = doc.paragraphs
    final_count = len(paras)
    print(f"\nDone. {initial_count} -> {final_count} paragraphs (deleted {initial_count - final_count})")

    doc.save(DOCX)
    print(f"Saved to {DOCX}")


if __name__ == "__main__":
    main()
